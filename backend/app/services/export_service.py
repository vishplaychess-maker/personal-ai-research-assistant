"""Report Export — document generation service.

Generates downloadable PDF (reportlab) and DOCX (python-docx) documents for
three export types:

  * ``research_report`` — a single markdown-ish report body (deep research
    reports surface as assistant messages in the chat, so the chat view sends
    the report text here).
  * ``chat``            — a conversation transcript (list of {role, content}).
  * ``knowledge_graph`` — the user's entity/relation graph as a table.

Everything is generated **in memory** (io.BytesIO) — no temporary files, no
filesystem writes, hence no path-traversal surface at all.

Security model:
  * ``type`` is validated against an enum — nothing else reaches the file.
  * Text content is sanitized: control characters stripped, and for PDF the
    text is escaped (reportlab Paragraph markup) before rendering.
  * Filenames are generated server-side from a sanitized title — the client
    never controls the filename, so no header-injection / traversal risk.
  * Payload size is bounded (MAX_* limits) to prevent memory exhaustion.
"""

import io
import logging
import re
from datetime import datetime, timezone
from enum import Enum
from typing import Any, Dict, List, Optional, Sequence

logger = logging.getLogger(__name__)

# ── Limits (bounded payloads, no memory exhaustion) ─────────
MAX_TITLE_LEN = 200
MAX_TEXT_CHARS = 500_000        # per report body / per chat message
MAX_CHAT_MESSAGES = 5_000
MAX_GRAPH_NODES = 20_000
MAX_GRAPH_EDGES = 50_000


class ExportTypeError(ValueError):
    """Raised when an invalid export ``type`` is requested."""


class ExportFormat(str, Enum):
    PDF = "pdf"
    DOCX = "docx"


class ExportType(str, Enum):
    RESEARCH_REPORT = "research_report"
    CHAT = "chat"
    KNOWLEDGE_GRAPH = "knowledge_graph"


# ── Sanitization helpers ────────────────────────────────────

# C0 control chars except \t \n \r; plus DEL and C1 range. These can corrupt
# binary formats or enable terminal/header tricks and never belong in a report.
_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f\x80-\x9f]")


def sanitize_text(text: Any, max_chars: int = MAX_TEXT_CHARS) -> str:
    """Coerce to ``str`` and strip control characters, bounded in length.

    All user/API-supplied text passes through this before reaching a
    document renderer.
    """
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    cleaned = _CONTROL_CHARS.sub("", text)
    if len(cleaned) > max_chars:
        cleaned = cleaned[:max_chars] + " …[truncated]"
    return cleaned


def sanitize_filename(title: Any) -> str:
    """Turn an arbitrary title into a safe single filename component.

    Keeps word characters (incl. unicode letters), spaces and hyphens;
    everything else collapses to ``_``. The result is always non-empty and
    cannot contain path separators, so path traversal is impossible.
    """
    name = sanitize_text(title, max_chars=MAX_TITLE_LEN).strip()
    name = re.sub(r"[^\w \-]+", "_", name, flags=re.UNICODE)
    name = re.sub(r"[\r\n\t]+", " ", name)
    name = re.sub(r"\s+", " ", name).strip(" ._")
    return (name or "export")[:MAX_TITLE_LEN]


def _strip_inline_markdown(text: str) -> str:
    """Remove common inline markdown for cleaner plain-text rendering."""
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)   # images -> alt text
    text = re.sub(r"\[([^\]]*)\]\([^)]*\)", r"\1", text)    # links -> text
    text = re.sub(r"`{1,3}([^`]*)`{1,3}", r"\1", text)      # inline code / fences
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)          # bold
    text = re.sub(r"(\*|_)([^*_]+)\1", r"\2", text)          # italic
    return text


def _markdown_blocks(text: str) -> List[Dict[str, Any]]:
    """Split sanitized markdown-ish text into renderable blocks.

    Returns a list of {kind, text, level} blocks where kind is one of
    ``heading`` | ``bullet`` | ``numbered`` | ``para``. Deliberately simple —
    this covers what the agent actually emits (headings, bullets, paras).
    """
    blocks: List[Dict[str, Any]] = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            blocks.append({"kind": "heading", "text": _strip_inline_markdown(m.group(2)), "level": len(m.group(1))})
            continue
        m = re.match(r"^[-*+]\s+(.*)$", stripped)
        if m:
            blocks.append({"kind": "bullet", "text": _strip_inline_markdown(m.group(1)), "level": 0})
            continue
        m = re.match(r"^(\d+)[.)]\s+(.*)$", stripped)
        if m:
            blocks.append({"kind": "numbered", "text": f"{m.group(1)}. {_strip_inline_markdown(m.group(2))}", "level": 0})
            continue
        blocks.append({"kind": "para", "text": _strip_inline_markdown(stripped), "level": 0})
    return blocks


# ── PDF (reportlab) ─────────────────────────────────────────


def _escape_pdf(text: str) -> str:
    """Escape reportlab Paragraph inline markup (XML-like)."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _pdf_flowables(blocks: Sequence[Dict[str, Any]]) -> List[Any]:
    """Convert markdown blocks into reportlab flowables."""
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, Spacer

    styles = getSampleStyleSheet()
    heading_styles = {
        1: styles["Heading1"],
        2: styles["Heading2"],
        3: styles["Heading3"],
    }

    flow: List[Any] = []
    for b in blocks:
        style = styles["BodyText"]
        prefix = ""
        if b["kind"] == "heading":
            style = heading_styles.get(min(b["level"], 3), styles["Heading3"])
        elif b["kind"] == "bullet":
            prefix = "\u2022&nbsp;&nbsp;"
        elif b["kind"] == "numbered":
            prefix = ""  # number already in text
        flow.append(Paragraph(prefix + _escape_pdf(b["text"]), style))
        flow.append(Spacer(1, 0.04 * inch))
    return flow


def generate_pdf(
    title: str,
    *,
    report_text: Optional[str] = None,
    chat_messages: Optional[List[Dict[str, str]]] = None,
    graph: Optional[Dict[str, List[Dict[str, Any]]]] = None,
) -> bytes:
    """Render a PDF to an in-memory buffer and return the bytes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
        Paragraph,
    )
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    flow: List[Any] = [
        Paragraph(_escape_pdf(sanitize_text(title, MAX_TITLE_LEN)), styles["Title"]),
        Paragraph(
            f"Exported from Thunder AI \u2014 {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
            styles["Normal"],
        ),
        Spacer(1, 0.25 * inch),
    ]

    if report_text is not None:
        flow.extend(_pdf_flowables(_markdown_blocks(sanitize_text(report_text))))

    if chat_messages is not None:
        for msg in chat_messages[:MAX_CHAT_MESSAGES]:
            role = sanitize_text(msg.get("role", "user"), 20).title() or "User"
            content = sanitize_text(msg.get("content", ""))
            when = sanitize_text(msg.get("created_at", ""), 40)
            header = f"{role}" + (f" \u00b7 {when}" if when else "")
            flow.append(Paragraph(f"<b>{_escape_pdf(header)}</b>", styles["Heading4"]))
            for b in _pdf_flowables(_markdown_blocks(content)):
                flow.append(b)
            flow.append(Spacer(1, 0.12 * inch))

    if graph is not None:
        nodes = graph.get("nodes", [])[:MAX_GRAPH_NODES]
        links = graph.get("links", [])[:MAX_GRAPH_EDGES]
        flow.append(Paragraph("Entities", styles["Heading2"]))
        node_names = [sanitize_text(n.get("name", ""), MAX_TITLE_LEN) for n in nodes]
        flow.append(Paragraph(_escape_pdf(", ".join(node_names) or "none"), styles["BodyText"]))
        flow.append(Spacer(1, 0.2 * inch))
        flow.append(Paragraph("Relations", styles["Heading2"]))
        rows = [["Source", "Relation", "Target", "Weight"]]
        for l in links:
            sid, tid = l.get("source"), l.get("target")
            src = node_names[sid] if isinstance(sid, int) and 0 <= sid < len(node_names) else sanitize_text(str(sid), 40)
            tgt = node_names[tid] if isinstance(tid, int) and 0 <= tid < len(node_names) else sanitize_text(str(tid), 40)
            rows.append([src, sanitize_text(l.get("relation", ""), 100), tgt, str(l.get("weight", 1))])
        table = Table(rows, colWidths=[1.8 * inch, 1.9 * inch, 1.8 * inch, 0.6 * inch], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#312e81")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
        ]))
        flow.append(table)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        title=sanitize_text(title, MAX_TITLE_LEN),
        author="Thunder AI",
    )
    doc.build(flow)
    return buf.getvalue()


# ── DOCX (python-docx) ──────────────────────────────────────


def generate_docx(
    title: str,
    *,
    report_text: Optional[str] = None,
    chat_messages: Optional[List[Dict[str, str]]] = None,
    graph: Optional[Dict[str, List[Dict[str, Any]]]] = None,
) -> bytes:
    """Render a DOCX to an in-memory buffer and return the bytes."""
    import docx
    from docx.shared import Pt, Inches

    document = docx.Document()
    document.core_properties.title = sanitize_text(title, MAX_TITLE_LEN)
    document.core_properties.author = "Thunder AI"

    document.add_heading(sanitize_text(title, MAX_TITLE_LEN), level=0)
    document.add_paragraph(
        f"Exported from Thunder AI — {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )

    def add_blocks(text: str) -> None:
        for b in _markdown_blocks(_strip_inline_markdown(sanitize_text(text))):
            if b["kind"] == "heading":
                document.add_heading(b["text"], level=min(b["level"], 3))
            elif b["kind"] == "bullet":
                document.add_paragraph(b["text"], style="List Bullet")
            elif b["kind"] == "numbered":
                document.add_paragraph(b["text"], style="List Number")
            else:
                document.add_paragraph(b["text"])

    if report_text is not None:
        add_blocks(report_text)

    if chat_messages is not None:
        for msg in chat_messages[:MAX_CHAT_MESSAGES]:
            role = sanitize_text(msg.get("role", "user"), 20).title() or "User"
            when = sanitize_text(msg.get("created_at", ""), 40)
            header = role + (f" · {when}" if when else "")
            document.add_heading(header, level=3)
            add_blocks(msg.get("content", ""))

    if graph is not None:
        nodes = graph.get("nodes", [])[:MAX_GRAPH_NODES]
        links = graph.get("links", [])[:MAX_GRAPH_EDGES]
        node_names = [sanitize_text(n.get("name", ""), MAX_TITLE_LEN) for n in nodes]
        document.add_heading("Entities", level=1)
        document.add_paragraph(", ".join(node_names) or "none")
        document.add_heading("Relations", level=1)
        table = document.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Source"; hdr[1].text = "Relation"; hdr[2].text = "Target"; hdr[3].text = "Weight"
        for l in links:
            row = table.add_row().cells
            sid, tid = l.get("source"), l.get("target")
            row[0].text = node_names[sid] if isinstance(sid, int) and 0 <= sid < len(node_names) else sanitize_text(str(sid), 40)
            row[1].text = sanitize_text(l.get("relation", ""), 100)
            row[2].text = node_names[tid] if isinstance(tid, int) and 0 <= tid < len(node_names) else sanitize_text(str(tid), 40)
            row[3].text = str(l.get("weight", 1))

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
